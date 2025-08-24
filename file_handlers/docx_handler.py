"""
Microsoft Word document (.docx) file handler.

This module provides the DocxHandler class for processing Microsoft Word (.docx) files
with fallback strategies for text extraction and content processing.
"""

import logging
import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional

from file_handlers.base_handler import BaseFileHandler
from utils.exceptions import FileProcessingError, ValidationError

# Try to import required libraries with fallbacks
try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    pypandoc = None
    PYPANDOC_AVAILABLE = False

try:
    import docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    docx = None
    PYTHON_DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocxHandler(BaseFileHandler):
    """
    Handler for Microsoft Word (.docx) files with multiple fallback options.

    This handler attempts to process .docx files using multiple strategies:
    1. Primary: python-docx library for structured content
    2. Alternative: pypandoc for text conversion
    3. Fallback: Basic text extraction

    The handler gracefully falls back to alternative methods when primary
    methods are unavailable or fail.
    """

    def __init__(self):
        """Initialize the DOCX file handler."""
        super().__init__()
        self.supported_extensions = ["docx"]
        self.handler_name = "Microsoft Word Handler"

        # Track which methods are available
        self._python_docx_available = PYTHON_DOCX_AVAILABLE
        self._pypandoc_available = PYPANDOC_AVAILABLE

        logger.info(
            f"DOCX Handler initialized - python-docx: {self._python_docx_available}, "
            f"pypandoc: {self._pypandoc_available}"
        )

    def can_handle(self, file_path: str) -> bool:
        """
        Check if this handler can process the given file.

        Args:
            file_path (str): Path to the file to check

        Returns:
            bool: True if this handler can process the file, False otherwise
        """
        if not file_path.lower().endswith(".docx"):
            return False

        # Check if at least one processing method is available
        return self._python_docx_available or self._pypandoc_available

    def extract_data(self, file_path: str) -> Dict[str, Any]:
        """
        Extract structured data from the DOCX file.

        Args:
            file_path (str): Path to the DOCX file to process

        Returns:
            Dict[str, Any]: Extracted data including text content, tables, and structure

        Raises:
            FileProcessingError: If the file cannot be processed by any method
        """
        if not os.path.exists(file_path):
            raise FileProcessingError(f"File not found: {file_path}")

        logger.info(f"Attempting to extract data from DOCX file: {file_path}")

        # Try each method in order of preference
        methods = [
            ("python-docx", self._extract_with_python_docx),
            ("pypandoc", self._extract_with_pypandoc),
        ]

        last_error = None
        for method_name, method_func in methods:
            try:
                if method_name == "python-docx" and not self._python_docx_available:
                    continue
                elif method_name == "pypandoc" and not self._pypandoc_available:
                    continue

                logger.info(f"Trying {method_name} method for {file_path}")
                data = method_func(file_path)
                logger.info(f"Successfully extracted data using {method_name} method")

                # Add metadata about extraction method
                data["extraction_metadata"] = {
                    "method_used": method_name,
                    "extraction_timestamp": str(Path(file_path).stat().st_mtime),
                    "file_path": file_path,
                    "file_size": os.path.getsize(file_path),
                }

                return data

            except Exception as e:
                logger.warning(f"{method_name} method failed for {file_path}: {str(e)}")
                last_error = e
                continue

        # If all methods failed
        error_msg = f"All extraction methods failed for {file_path}"
        if last_error:
            error_msg += f". Last error: {str(last_error)}"

        logger.error(error_msg)
        raise FileProcessingError(error_msg)

    def _extract_with_python_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract data using python-docx library.

        Args:
            file_path (str): Path to the DOCX file

        Returns:
            Dict[str, Any]: Extracted structured data
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise FileProcessingError("python-docx library is not available")

        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text.strip(),
                        "style": para.style.name if para.style else "Normal"
                    })

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = {
                    "headers": [],
                    "rows": []
                }
                
                # Get headers from first row
                if table.rows:
                    first_row = table.rows[0]
                    table_data["headers"] = [cell.text.strip() for cell in first_row.cells]
                    
                    # Get data rows
                    for row in table.rows[1:]:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data["rows"].append(row_data)
                
                tables.append(table_data)

            # Extract raw content for text-based parsing
            raw_content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            # Extract sections based on headings
            sections = []
            current_section = None
            for para in doc.paragraphs:
                if para.style and ('Heading' in para.style.name or 'Title' in para.style.name):
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        "title": para.text.strip(),
                        "level": para.style.name,
                        "content": ""
                    }
                elif current_section and para.text.strip():
                    current_section["content"] += para.text.strip() + "\n"

            if current_section:
                sections.append(current_section)

            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "sections": sections,
                "raw_content": raw_content,
                "document_properties": {
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "section_count": len(sections)
                }
            }

        except Exception as e:
            raise FileProcessingError(f"python-docx extraction failed: {str(e)}") from e

    def _extract_with_pypandoc(self, file_path: str) -> Dict[str, Any]:
        """
        Extract data using pypandoc for text conversion.

        Args:
            file_path (str): Path to the DOCX file

        Returns:
            Dict[str, Any]: Extracted text data
        """
        if not PYPANDOC_AVAILABLE:
            raise FileProcessingError("pypandoc library is not available")

        try:
            # Convert to markdown first to preserve structure
            markdown_content = pypandoc.convert_file(
                file_path, 
                'markdown',
                extra_args=['--standalone', '--wrap=none']
            )

            # Convert to plain text as fallback
            plain_content = pypandoc.convert_file(
                file_path,
                'plain',
                extra_args=['--wrap=none']
            )

            # Basic section parsing from markdown
            sections = []
            current_section = None
            
            for line in markdown_content.split('\n'):
                line = line.strip()
                if line.startswith('#'):
                    if current_section:
                        sections.append(current_section)
                    
                    # Count heading level
                    level = len(line) - len(line.lstrip('#'))
                    title = line.lstrip('#').strip()
                    
                    current_section = {
                        "title": title,
                        "level": f"Heading {level}",
                        "content": ""
                    }
                elif current_section and line:
                    current_section["content"] += line + "\n"

            if current_section:
                sections.append(current_section)

            # Basic table extraction from markdown
            tables = []
            lines = markdown_content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
                    # Found a table
                    headers = [cell.strip() for cell in line.split('|')[1:-1]]
                    i += 2  # Skip header separator
                    
                    rows = []
                    while i < len(lines) and '|' in lines[i]:
                        row_data = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                        if len(row_data) == len(headers):
                            rows.append(row_data)
                        i += 1
                    
                    tables.append({
                        "headers": headers,
                        "rows": rows
                    })
                else:
                    i += 1

            return {
                "paragraphs": [{"text": para, "style": "Normal"} 
                              for para in plain_content.split('\n\n') if para.strip()],
                "tables": tables,
                "sections": sections,
                "raw_content": plain_content,
                "markdown_content": markdown_content,
                "document_properties": {
                    "paragraph_count": len([p for p in plain_content.split('\n\n') if p.strip()]),
                    "table_count": len(tables),
                    "section_count": len(sections)
                }
            }

        except Exception as e:
            raise FileProcessingError(f"pypandoc extraction failed: {str(e)}") from e

    def validate_structure(self, file_path: str) -> Dict[str, Any]:
        """
        Validate the DOCX file structure and content.

        Args:
            file_path (str): Path to the DOCX file to validate

        Returns:
            Dict[str, Any]: Validation result with success status and messages
        """
        result = {"is_valid": True, "errors": [], "warnings": []}

        try:
            # Basic file existence and extension check
            if not os.path.exists(file_path):
                result["errors"].append(f"File does not exist: {file_path}")
                result["is_valid"] = False
                return result

            if not file_path.lower().endswith(".docx"):
                result["errors"].append(f"File does not have .docx extension: {file_path}")
                result["is_valid"] = False
                return result

            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                result["errors"].append("DOCX file is empty")
                result["is_valid"] = False
                return result
            elif file_size < 1024:  # Less than 1KB is suspicious for DOCX
                result["warnings"].append("DOCX file is unusually small (< 1KB)")

            # Check if any processing method is available
            if not (self._python_docx_available or self._pypandoc_available):
                result["errors"].append("No DOCX processing methods are available")
                result["is_valid"] = False
                return result

            # Try to extract basic data to validate file integrity
            try:
                data = self.extract_data(file_path)

                # Validate extracted data structure
                if not data.get("raw_content", "").strip():
                    result["warnings"].append("No text content found in DOCX file")

                if not data.get("paragraphs") and not data.get("tables"):
                    result["warnings"].append("No structured content found in DOCX file")

            except Exception as e:
                result["errors"].append(f"Failed to validate DOCX file content: {str(e)}")
                result["is_valid"] = False

        except Exception as e:
            result["errors"].append(f"Validation error: {str(e)}")
            result["is_valid"] = False

        return result